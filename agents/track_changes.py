"""
Generación de .docx con control de cambios y comentarios al margen.

Agentes 1-3 (FORMA, ESTILO JUDICIAL, COHERENCIA NARRATIVA):
  Inserta <w:del> (rojo tachado) + <w:ins> (verde subrayado) donde se encuentre
  el texto de ubicacion. Si no se encuentra, cae a comentario.

Agentes 4-5 (FONDO ARGUMENTATIVO, NORMATIVO):
  Inserta comentario al margen con error, ubicacion y correccion.
"""

import copy
import io
import re
import zipfile
from datetime import datetime, timezone
from lxml import etree

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"

AGENTES_TRACK = {"FORMA", "ESTILO JUDICIAL", "COHERENCIA NARRATIVA"}


def _now() -> str:
    return datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")


# ── Mapa de offsets LT ───────────────────────────────────────────────────────

def _build_para_offset_map(doc: Document) -> list[tuple[int, int, int]]:
    """
    Reconstruye el mismo texto plano que main.py extrae del .docx:
        texto = "\\n".join(p.text for p in doc.paragraphs if p.text.strip())

    Retorna una lista ordenada de (offset_inicio, offset_fin, para_idx_en_doc)
    para cada párrafo no vacío. Permite mapear un offset LT al párrafo exacto
    sin búsqueda aproximada.
    """
    result: list[tuple[int, int, int]] = []
    pos = 0
    for i, para in enumerate(doc.paragraphs):
        if not para.text.strip():
            continue
        text = para.text
        result.append((pos, pos + len(text), i))
        pos += len(text) + 1  # +1 por el "\\n" del join en main.py
    return result


def _find_paragraph_by_offset(
    offset_map: list[tuple[int, int, int]], lt_offset: int
) -> int | None:
    """
    Retorna el índice docx del párrafo que contiene la posición lt_offset,
    o None si el offset cae fuera del rango del documento.
    La lista está ordenada, así que se detiene en el primer elemento posterior.
    """
    for start, end, para_idx in offset_map:
        if start <= lt_offset < end:
            return para_idx
        if start > lt_offset:
            break
    return None


# ── Búsqueda de texto ─────────────────────────────────────────────────────────

def _find_paragraph(doc: Document, texto: str) -> int | None:
    """
    Busca el párrafo que contiene 'texto'.
    1. Coincidencia exacta de subcadena.
    2. Coincidencia de las primeras palabras significativas (maneja
       truncamiento, comillas o diferencias menores del AI).
    Nunca retorna el último párrafo como fallback — eso lo decide el llamador.
    """
    needle = texto.lower().strip()
    if not needle:
        return None

    # 1. Coincidencia exacta
    for i, para in enumerate(doc.paragraphs):
        if needle in para.text.lower():
            return i

    # 2. Fuzzy: palabras significativas (>3 chars) de los primeros 80 caracteres
    palabras = [w for w in re.sub(r"[^\w\s]", " ", needle[:80]).split() if len(w) > 3][:6]
    if len(palabras) >= 2:
        best_idx, best_hits = None, 0
        for i, para in enumerate(doc.paragraphs):
            pt = para.text.lower()
            hits = sum(1 for w in palabras if w in pt)
            if hits > best_hits:
                best_hits, best_idx = hits, i
        # Exige al menos la mitad de las palabras clave
        if best_idx is not None and best_hits >= max(2, len(palabras) // 2):
            return best_idx

    return None


def _run_map(para_el) -> list[tuple[int, int, object, str]]:
    """
    Retorna [(start, end, run_el, text)] para TODOS los <w:r> del párrafo,
    recursivamente a cualquier profundidad de anidamiento.
    Excluye runs dentro de <w:del> (texto eliminado no visible).
    """
    result, pos = [], 0
    seen = set()
    for r in para_el.iter(f"{{{W}}}r"):
        if id(r) in seen:
            continue
        seen.add(id(r))
        # Subir hasta para_el buscando un ancestro <w:del>
        ancestor, in_del = r.getparent(), False
        while ancestor is not None and ancestor is not para_el:
            if etree.QName(ancestor).localname == "del":
                in_del = True
                break
            ancestor = ancestor.getparent()
        if in_del:
            continue
        # Recopilar todos los <w:t> del run (usualmente uno, a veces varios)
        texts = [t.text or "" for t in r.findall(f"{{{W}}}t")]
        text = "".join(texts)
        result.append((pos, pos + len(text), r, text))
        pos += len(text)
    return result


# ── Track changes (w:del + w:ins) ────────────────────────────────────────────

def _make_rpr(orig_rpr, color: str, extra: list[str]) -> object:
    rpr = OxmlElement("w:rPr")
    if orig_rpr is not None:
        skip = {"color", "strike", "u", "rStyle"}
        for child in orig_rpr:
            if etree.QName(child).localname not in skip:
                rpr.append(copy.deepcopy(child))
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    rpr.append(c)
    for tag in extra:
        el = OxmlElement(tag)
        if tag == "w:u":
            el.set(qn("w:val"), "single")
        rpr.append(el)
    return rpr


def _make_t(tag: str, text: str) -> object:
    el = OxmlElement(tag)
    el.text = text
    if text and (text[0] == " " or text[-1] == " "):
        el.set(XML_SPACE, "preserve")
    return el


def insertar_track_change(doc: Document, para_idx: int, ubicacion: str,
                          correccion: str, autor: str, rev_id: int) -> bool:
    """
    Busca ubicacion en el párrafo y reemplaza con w:del + w:ins.
    Retorna True si tuvo éxito.
    """
    para_el = doc.paragraphs[para_idx]._p
    rmap = _run_map(para_el)
    if not rmap:
        return False

    combined = "".join(t for _, _, _, t in rmap)
    pos = combined.lower().find(ubicacion.lower().strip())
    if pos == -1:
        return False

    end = pos + len(ubicacion)
    texto_real = combined[pos:end]
    affected = [(s, e, r, t) for s, e, r, t in rmap if s < end and e > pos]
    if not affected:
        return False

    now = _now()
    first_r = affected[0][2]
    orig_rpr = first_r.find(f"{{{W}}}rPr")

    # <w:del>
    del_el = OxmlElement("w:del")
    del_el.set(qn("w:id"), str(rev_id))
    del_el.set(qn("w:author"), autor)
    del_el.set(qn("w:date"), now)
    del_r = OxmlElement("w:r")
    del_r.append(_make_rpr(orig_rpr, "FF0000", ["w:strike"]))
    del_r.append(_make_t("w:delText", texto_real))
    del_el.append(del_r)

    # <w:ins>
    ins_el = OxmlElement("w:ins")
    ins_el.set(qn("w:id"), str(rev_id + 1))
    ins_el.set(qn("w:author"), autor)
    ins_el.set(qn("w:date"), now)
    ins_r = OxmlElement("w:r")
    ins_r.append(_make_rpr(orig_rpr, "00B050", ["w:u"]))
    ins_r.append(_make_t("w:t", correccion))
    ins_el.append(ins_r)

    # Posición de inserción
    parent = first_r.getparent()
    insert_idx = list(parent).index(first_r)

    first_s = affected[0][0]
    last_e = affected[-1][1]
    prefix = combined[first_s:pos]
    suffix = combined[end:last_e]

    # Eliminar runs afectados
    for _, _, r_el, _ in affected:
        if r_el.getparent() is not None:
            r_el.getparent().remove(r_el)

    idx = insert_idx

    if prefix:
        pre_r = copy.deepcopy(first_r)
        t_el = pre_r.find(f"{{{W}}}t")
        if t_el is None:
            t_el = OxmlElement("w:t")
            pre_r.append(t_el)
        t_el.text = prefix
        if prefix[0] == " " or prefix[-1] == " ":
            t_el.set(XML_SPACE, "preserve")
        parent.insert(idx, pre_r)
        idx += 1

    parent.insert(idx, del_el);  idx += 1
    parent.insert(idx, ins_el);  idx += 1

    if suffix:
        suf_r = copy.deepcopy(affected[-1][2])
        t_el = suf_r.find(f"{{{W}}}t")
        if t_el is None:
            t_el = OxmlElement("w:t")
            suf_r.append(t_el)
        t_el.text = suffix
        if suffix[0] == " " or suffix[-1] == " ":
            t_el.set(XML_SPACE, "preserve")
        parent.insert(idx, suf_r)

    return True


# ── Comentarios al margen ─────────────────────────────────────────────────────

def _add_comment_ref(para_el, comment_id: int, ubicacion: str = "") -> None:
    """
    Ancla commentRangeStart / commentRangeEnd / commentReference al run
    específico que contiene 'ubicacion'.

    Estructura OOXML requerida:
        <w:commentRangeStart w:id="N"/>
        <w:r>...<w:t>texto con el error</w:t></w:r>
        <w:commentRangeEnd w:id="N"/>
        <w:r><w:rPr><w:rStyle w:val="CommentReference"/></w:rPr>
             <w:commentReference w:id="N"/></w:r>

    commentRangeEnd y commentReference DEBEN ser adyacentes y ubicados
    justo después del run anotado — de lo contrario Word los agrupa al final.
    """
    def _make_start():
        el = OxmlElement("w:commentRangeStart")
        el.set(qn("w:id"), str(comment_id))
        return el

    def _make_end():
        el = OxmlElement("w:commentRangeEnd")
        el.set(qn("w:id"), str(comment_id))
        return el

    def _make_ref_run():
        r = OxmlElement("w:r")
        rpr = OxmlElement("w:rPr")
        sty = OxmlElement("w:rStyle")
        sty.set(qn("w:val"), "CommentReference")
        rpr.append(sty)
        r.append(rpr)
        cref = OxmlElement("w:commentReference")
        cref.set(qn("w:id"), str(comment_id))
        r.append(cref)
        return r

    # ── Intentar anclar al run exacto que contiene ubicacion ─────────────────
    target_run = None
    if ubicacion:
        needle = ubicacion.lower().strip()
        rmap = _run_map(para_el)
        combined = "".join(t for _, _, _, t in rmap)
        pos = combined.lower().find(needle)
        if pos != -1:
            # Run donde COMIENZA el fragmento anotado
            for s, e, r_el, _ in rmap:
                if s <= pos < e:
                    target_run = r_el
                    break

    if target_run is not None:
        # Subir hasta encontrar el hijo directo del w:p.
        # Si target_run está dentro de w:hyperlink u otro contenedor,
        # insertar dentro de ese contenedor haría que Word desplace
        # los marcadores al final del documento.
        anchor = target_run
        while etree.QName(anchor.getparent()).localname != "p":
            anchor = anchor.getparent()
            if anchor.getparent() is None:
                break  # seguridad: no debería ocurrir

        parent = anchor.getparent()  # siempre es el w:p
        idx = list(parent).index(anchor)
        parent.insert(idx, _make_start())      # idx   → commentRangeStart
        # anchor ahora está en idx + 1
        parent.insert(idx + 2, _make_end())    # idx+2 → commentRangeEnd
        parent.insert(idx + 3, _make_ref_run())# idx+3 → commentReference
        return

    # ── Fallback: anclar al primer run disponible del párrafo ────────────────
    children = list(para_el)
    first_run_idx = next(
        (i for i, c in enumerate(children)
         if etree.QName(c).localname in ("r", "hyperlink", "ins", "del")),
        None,
    )
    if first_run_idx is None:
        # Párrafo vacío: solo appender al final
        para_el.append(_make_start())
        para_el.append(_make_end())
        para_el.append(_make_ref_run())
        return

    para_el.insert(first_run_idx, _make_start())
    # Recalcular índice tras insert
    children = list(para_el)
    last_run_idx = max(
        i for i, c in enumerate(children)
        if etree.QName(c).localname in ("r", "hyperlink", "ins", "del")
    )
    para_el.insert(last_run_idx + 1, _make_end())
    para_el.insert(last_run_idx + 2, _make_ref_run())


def _build_comment_xml(comment_id: int, autor: str, texto: str, fecha: str) -> str:
    def esc(s: str) -> str:
        return (s.replace("&", "&amp;").replace("<", "&lt;")
                 .replace(">", "&gt;").replace('"', "&quot;"))

    paras = texto.split("\n")
    paras_xml = ""
    for i, p in enumerate(paras):
        paras_xml += f'<w:p><w:pPr><w:pStyle w:val="CommentText"/></w:pPr>'
        if i == 0:
            paras_xml += ('<w:r><w:rPr><w:rStyle w:val="CommentReference"/></w:rPr>'
                          '<w:annotationRef/></w:r>')
        paras_xml += f'<w:r><w:t xml:space="preserve">{esc(p)}</w:t></w:r></w:p>'

    return (f'<w:comment w:id="{comment_id}" w:author="{esc(autor)}" '
            f'w:date="{fecha}" w:initials="DIA">{paras_xml}</w:comment>')


def _inject_comments(docx_bytes: bytes, comment_xmls: list[str]) -> bytes:
    """Inyecta word/comments.xml en el zip del .docx."""
    inner = "\n".join(comment_xmls)
    comments_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f"{inner}"
        "</w:comments>"
    )

    src = zipfile.ZipFile(io.BytesIO(docx_bytes), "r")
    out_buf = io.BytesIO()

    with zipfile.ZipFile(out_buf, "w", zipfile.ZIP_DEFLATED) as out:
        for item in src.infolist():
            # Omitir comments.xml existente — lo reescribimos al final
            if item.filename == "word/comments.xml":
                continue

            data = src.read(item.filename)

            if item.filename == "word/_rels/document.xml.rels":
                text = data.decode("utf-8")
                if "comments" not in text.lower():
                    text = text.replace(
                        "</Relationships>",
                        '<Relationship Id="rIdComments" '
                        'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments" '
                        'Target="comments.xml"/>\n</Relationships>',
                    )
                out.writestr(item, text.encode("utf-8"))

            elif item.filename == "[Content_Types].xml":
                text = data.decode("utf-8")
                if "comments" not in text.lower():
                    text = text.replace(
                        "</Types>",
                        '<Override PartName="/word/comments.xml" '
                        'ContentType="application/vnd.openxmlformats-officedocument'
                        '.wordprocessingml.comments+xml"/>\n</Types>',
                    )
                out.writestr(item, text.encode("utf-8"))

            else:
                out.writestr(item, data)

        out.writestr("word/comments.xml", comments_xml.encode("utf-8"))

    src.close()
    return out_buf.getvalue()


# ── Corrección directa (sin track changes) ───────────────────────────────────

def _norm_ws(s: str) -> str:
    """Normaliza espacios tipográficos para matching robusto."""
    return re.sub(r"[     \t]+", " ", s)


def _reemplazar_texto_en_parrafo(para_el, ubicacion: str, correccion: str) -> bool:
    """
    Sustituye ubicacion por correccion directamente en los <w:t> del párrafo.
    No inserta <w:del>/<w:ins> — el resultado es texto limpio sin marcas de revisión.
    Retorna True si el reemplazo tuvo éxito.
    """
    rmap = _run_map(para_el)
    if not rmap:
        return False

    combined = "".join(t for _, _, _, t in rmap)
    needle = _norm_ws(ubicacion).lower().strip()
    haystack = _norm_ws(combined).lower()

    pos = haystack.find(needle)
    if pos == -1:
        return False

    end_pos = pos + len(needle)
    affected = [(s, e, r, t) for s, e, r, t in rmap if s < end_pos and e > pos]
    if not affected:
        return False

    # Texto que precede y sigue al match dentro del rango afectado
    prefix = combined[affected[0][0] : pos]
    suffix = combined[end_pos : affected[-1][1]]

    # Consolidar todo en el primer run afectado
    first_r = affected[0][2]
    t_el = first_r.find(f"{{{W}}}t")
    if t_el is None:
        t_el = OxmlElement("w:t")
        first_r.append(t_el)
    # Eliminar <w:t> adicionales del primer run si los hubiera
    for extra in first_r.findall(f"{{{W}}}t")[1:]:
        first_r.remove(extra)

    t_el.text = prefix + correccion + suffix
    if t_el.text and (t_el.text[0] == " " or t_el.text[-1] == " "):
        t_el.set(XML_SPACE, "preserve")

    # Eliminar los runs sobrantes absorbidos en el primero
    for _, _, r_el, _ in affected[1:]:
        parent = r_el.getparent()
        if parent is not None:
            parent.remove(r_el)

    return True


def aplicar_correcciones_zip(docx_bytes: bytes, hallazgos: list) -> bytes:
    """
    Aplica las correcciones aceptadas directamente al texto del .docx.

    Diferencias clave respecto a generar_documento_revisado():
      - Sin track changes: el texto queda limpio, listo para usar.
      - Corrección ZIP quirúrgica: solo se reescribe word/document.xml
        dentro del ZIP. Imágenes, estilos, headers, footers, rels y temas
        salen intactos del archivo original — sin pasar por doc.save().

    Restricciones idénticas a BUG 4: descarta hallazgos cuya ubicacion
    no existe literalmente en el documento. Descarta correcciones > 15 palabras
    (son reescrituras, no sustituciones puntuales).
    """
    from lxml import etree as letree

    doc = Document(io.BytesIO(docx_bytes))
    offset_map = _build_para_offset_map(doc)
    texto_doc = "\n".join(p.text for p in doc.paragraphs if p.text.strip()).lower()

    for h in hallazgos:
        ubicacion  = (h.ubicacion  or "").strip()
        correccion = (h.correccion or "").strip()

        if not ubicacion or not correccion:
            continue
        if ubicacion.lower() == correccion.lower():
            continue
        if ubicacion.lower() not in texto_doc:
            continue
        if len(correccion.split()) > 15:
            continue

        lt_offset = getattr(h, "lt_offset", None)
        if lt_offset is not None:
            para_idx = _find_paragraph_by_offset(offset_map, lt_offset)
        else:
            para_idx = _find_paragraph(doc, ubicacion) if ubicacion else None

        if para_idx is None:
            continue

        _reemplazar_texto_en_parrafo(doc.paragraphs[para_idx]._p, ubicacion, correccion)

    # Serializar el árbol lxml modificado (los _p son referencias directas al árbol)
    new_doc_xml = letree.tostring(
        doc.element,
        xml_declaration=True,
        encoding="UTF-8",
        standalone=True,
    )

    # Reensamblar ZIP: solo word/document.xml cambia, todo lo demás del original
    out_buf = io.BytesIO()
    with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as src:
        with zipfile.ZipFile(out_buf, "w", zipfile.ZIP_DEFLATED) as out:
            for item in src.infolist():
                if item.filename == "word/document.xml":
                    out.writestr(item, new_doc_xml)
                else:
                    out.writestr(item, src.read(item.filename))

    return out_buf.getvalue()


# ── Función principal ─────────────────────────────────────────────────────────

def generar_documento_revisado(docx_bytes: bytes, hallazgos: list) -> bytes:
    """
    Recibe bytes de .docx + lista de Hallazgo.
    Retorna bytes del .docx con track changes y comentarios insertados.

    Estrategia de localización de párrafo (por orden de prioridad):
      1. Hallazgos de LanguageTool con lt_offset → mapa de offsets exacto
         (misma lógica de extracción de texto que main.py)
      2. Resto de hallazgos → búsqueda aproximada por texto (fuzzy)
      3. Si ninguna encuentra el párrafo → comentario/track change se omite
         (nunca anclar al párrafo equivocado)
    """
    doc = Document(io.BytesIO(docx_bytes))
    rev_id = 200
    comment_id = 1
    comment_xmls: list[str] = []
    fecha = _now()

    # Mapa LT offset → párrafo (construido una sola vez)
    offset_map = _build_para_offset_map(doc)

    # Texto plano del documento para verificar existencia literal de ubicaciones
    _texto_doc = "\n".join(p.text for p in doc.paragraphs if p.text.strip()).lower()

    for h in hallazgos:
        ubicacion = (h.ubicacion or "").strip()
        correccion = (h.correccion or "").strip()
        error     = (h.error or "").strip()
        modulo    = (h.modulo or "").strip()
        agente    = h.agente
        severidad = h.severidad
        lt_offset = getattr(h, "lt_offset", None)

        if not error and not ubicacion:
            continue

        # Descartar hallazgos cuya ubicacion no existe literalmente en el documento.
        # Previene que alucinaciones del LLM produzcan comentarios en el párrafo
        # equivocado o track changes con texto que nunca estuvo en el documento.
        if ubicacion and ubicacion.lower() not in _texto_doc:
            continue

        autor = f"DISCIPLINAR[IA] — {agente} [{modulo}]"

        # ── Localizar párrafo objetivo (una sola vez por hallazgo) ───────────
        if lt_offset is not None:
            # LT: posición exacta en el texto plano → párrafo determinista
            para_idx = _find_paragraph_by_offset(offset_map, lt_offset)
        else:
            # LLM / folios: búsqueda aproximada por texto
            para_idx = _find_paragraph(doc, ubicacion) if ubicacion else None
        # Si no se encuentra el párrafo, el hallazgo se omite completamente.
        # Un comentario o track change en el párrafo equivocado es peor
        # que no insertarlo.
        if para_idx is None:
            continue

        inserted_track = False

        # ── Agentes 1-3: intentar track change ──────────────────────────────
        # Correcciones de más de 15 palabras son reescrituras de párrafo completo —
        # insertarlas como w:del/w:ins produce texto fabricado. Van a comentario.
        if (agente in AGENTES_TRACK
                and ubicacion and correccion
                and ubicacion.lower() != correccion.lower()
                and len(correccion.split()) <= 15):
            try:
                inserted_track = insertar_track_change(
                    doc, para_idx, ubicacion, correccion, autor, rev_id
                )
                if inserted_track:
                    rev_id += 2
            except Exception:
                inserted_track = False

        # ── Comentario al margen ─────────────────────────────────────────────
        # Agentes 4-5 siempre comentario.
        # Agentes 1-3: comentario si el track change no pudo insertarse.
        if not inserted_track:
            sev_tag = {"alta": "[ALTA]", "media": "[MEDIA]", "baja": "[BAJA]"}.get(
                severidad, ""
            )
            lineas = [f"{sev_tag} {modulo}: {error}"]
            if ubicacion:
                lineas.append(f'Texto: "{ubicacion[:120]}"')
            if correccion:
                lineas.append(f"Sugerencia: {correccion}")
            texto_comentario = "\n".join(lineas)

            try:
                _add_comment_ref(doc.paragraphs[para_idx]._p, comment_id, ubicacion)
                comment_xmls.append(
                    _build_comment_xml(comment_id, autor, texto_comentario, fecha)
                )
                comment_id += 1
            except Exception:
                pass  # nunca interrumpir la exportación por un comentario fallido

    # Guardar documento con track changes
    buf = io.BytesIO()
    doc.save(buf)
    result = buf.getvalue()

    # Inyectar comentarios si los hay
    if comment_xmls:
        result = _inject_comments(result, comment_xmls)

    return result
